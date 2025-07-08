import os
import markdown
from docx import Document as DocxDocument
from markdown2 import markdown as markdown2_html
from rich.console import Console
from rich.markdown import Markdown
import ray
import json

# For RTF conversion
from pyth.plugins.rtf15.reader import Rtf15Reader
from pyth.plugins.plaintext.writer import PlaintextWriter
from pyth.plugins.rtf15.writer import Rtf15Writer
from io import StringIO, BytesIO

# Load the markdown file
md_path = "/Users/rickhuang/Library/CloudStorage/OneDrive-Personal/workspace-mac/llama_index/data/eg_sample_text_doc.md"
with open(md_path, 'r', encoding='utf-8') as f:
    md_content = f.read()

# Convert markdown to HTML
html_content = markdown2_html(md_content)

# Convert HTML to DOCX
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

def html_to_docx(html, docx_path):
    doc = DocxDocument()
    soup = BeautifulSoup(html, 'html.parser')
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'ul', 'ol', 'pre', 'code']):
        if elem.name.startswith('h'):
            doc.add_heading(elem.get_text(), int(elem.name[1]))
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text())
        elif elem.name == 'li':
            doc.add_paragraph(f"- {elem.get_text()}")
        elif elem.name == 'pre' or elem.name == 'code':
            doc.add_paragraph(elem.get_text(), style='Intense Quote')
    doc.save(docx_path)

# Save as DOCX
converted_docx_path = md_path.replace('.md', '_converted.docx')
html_to_docx(html_content, converted_docx_path)

# Save as RTF (optional, using python-docx-rtf or pyth)
def docx_to_rtf(docx_path, rtf_path):
    # This is a placeholder: python-docx does not support RTF export directly.
    # You can use LibreOffice in headless mode or other tools for robust conversion.
    pass

# Save as TXT (plain text)
converted_txt_path = md_path.replace('.md', '_converted.txt')
with open(converted_txt_path, 'w', encoding='utf-8') as f:
    f.write(md_content)

print(f"Converted markdown to: {converted_docx_path} and {converted_txt_path}")

# Optionally, display the markdown in the terminal with formatting
console = Console()
console.print(Markdown(md_content))

# You can now integrate this with your agent logic from starter_agent.py as needed.

# --- Ray Distributed Markdown Conversion ---
ray.init(ignore_reinit_error=True)

@ray.remote
def convert_md_to_docx_ray(md_path):
    from markdown2 import markdown as markdown2_html
    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    html_content = markdown2_html(md_content)
    doc = DocxDocument()
    soup = BeautifulSoup(html_content, 'html.parser')
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'ul', 'ol', 'pre', 'code']):
        if elem.name.startswith('h'):
            doc.add_heading(elem.get_text(), int(elem.name[1]))
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text())
        elif elem.name == 'li':
            doc.add_paragraph(f"- {elem.get_text()}")
        elif elem.name == 'pre' or elem.name == 'code':
            doc.add_paragraph(elem.get_text(), style='Intense Quote')
    out_path = md_path.replace('.md', '_ray_converted.docx')
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    # Distributed conversion example
    md_files = [
        "/Users/rickhuang/Library/CloudStorage/OneDrive-Personal/workspace-mac/llama_index/data/eg_sample_text_doc.md",
        # Add more markdown files here if needed
    ]
    futures = [convert_md_to_docx_ray.remote(md) for md in md_files]
    results = ray.get(futures)
    print("Ray converted files:", results)

    # --- Ray RLlib Reinforcement Learning Example ---
    from ray import tune
    def train_cartpole(config):
        import gym
        env = gym.make("CartPole-v1")
        obs = env.reset()
        total_reward = 0
        for _ in range(200):
            action = env.action_space.sample()
            obs, reward, done, info = env.step(action)
            total_reward += reward
            if done:
                break
        tune.report(reward=total_reward)

    analysis = tune.run(
        train_cartpole,
        config={},
        num_samples=2,
        resources_per_trial={"cpu": 1}
    )
    print("Best RL trial:", analysis.get_best_trial("reward", mode="max"))

# --- Langraph Integration: Load and use langraph.json for workflow configuration ---
LANGRAPH_CONFIG_PATH = os.path.join(os.path.dirname(__file__), 'langraph.json')
if os.path.exists(LANGRAPH_CONFIG_PATH):
    with open(LANGRAPH_CONFIG_PATH, 'r') as f:
        langraph_config = json.load(f)
    print("[Langraph] Loaded configuration:")
    print(json.dumps(langraph_config, indent=2))
else:
    langraph_config = None
    print("[Langraph] No configuration file found.")

# Example: Use Langraph config to control workflow (stub)
def run_langraph_workflow(input_data, config=langraph_config):
    if not config:
        print("No Langraph config loaded. Skipping advanced workflow.")
        return None
    # Example: Use config to set max agents/tools, enable visualization, etc.
    print(f"[Langraph] Running workflow with up to {config['default_settings']['max_agents']} agents and {config['default_settings']['max_tools']} tools.")
    # ... insert Langraph workflow logic here ...
    return "[Langraph] Workflow completed (stub)"

# Example usage
if __name__ == "__main__":
    # ...existing code...
    run_langraph_workflow(input_data={"example": "data"})
